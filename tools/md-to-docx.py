"""
제출 문서 초안(markdown) → Word(.docx) 변환

공식 양식(hwpx)의 구조를 그대로 따른다. 제목, 팀명·구성원 표, 번호 붙은 항목 순서가
양식과 같아야 한글로 옮겨 담을 때 대조가 쉽고, 그대로 PDF로 내보내도 형태가 맞는다.

    python tools/md-to-docx.py

팀명·구성원은 TEAM 에서 읽는다. 값을 채운 뒤 다시 실행하면 반영된다.
"""

from __future__ import annotations

import pathlib
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ── 여기만 채우면 된다 ──────────────────────────────────────────
TEAM = {
    "팀명": "",           # 대회 사이트에 등록된 팀명과 정확히 같게
    "구성원 성명": "",     # 팀장, 팀원 순
}

BODY_FONT = "맑은 고딕"
MONO_FONT = "D2Coding"  # 없으면 Word 가 대체한다
INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x5A, 0x5A, 0x5A)

DOCS = [
    {
        "src": "docs/제출-기획서.md",
        "out": "제출/(첨부1) 2026 금융 AI Challenge 기획서.docx",
        "title": "2026 금융 AI Challenge 기획서",
        "tag": "첨부 1",
    },
    {
        "src": "docs/제출-기능명세서.md",
        "out": "제출/(첨부2) 2026 금융 AI Challenge 기능 명세서.docx",
        "title": "2026 금융 AI Challenge 기능 명세서",
        "tag": "첨부 2",
    },
]


# ── 문서 기본 설정 ────────────────────────────────────────────
def setup(doc: Document) -> None:
    s = doc.sections[0]
    s.page_width, s.page_height = Cm(21.0), Cm(29.7)  # A4
    s.left_margin = s.right_margin = Cm(2.0)
    s.top_margin = s.bottom_margin = Cm(2.0)

    st = doc.styles["Normal"]
    st.font.name = BODY_FONT
    st.font.size = Pt(10)
    st.font.color.rgb = INK
    st.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    pf = st.paragraph_format
    pf.space_before, pf.space_after = Pt(0), Pt(4)
    pf.line_spacing = 1.35


def run_font(run, name: str = BODY_FONT, size: float | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)


# ── 인라인 서식: **굵게**, `코드`, *기울임* ────────────────────
INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\*[^*\s][^*]*?\*)")


def add_inline(par, text: str, size: float | None = None) -> None:
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)  # 링크는 글자만
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = par.add_run(part[2:-2])
            r.bold = True
            run_font(r, BODY_FONT, size)
        elif part.startswith("`") and part.endswith("`"):
            r = par.add_run(part[1:-1])
            run_font(r, MONO_FONT, (size or 10) - 0.5)
        elif part.startswith("*") and part.endswith("*"):
            r = par.add_run(part[1:-1])
            r.italic = True
            run_font(r, BODY_FONT, size)
        else:
            r = par.add_run(part)
            run_font(r, BODY_FONT, size)


# ── 블록 요소 ────────────────────────────────────────────────
def heading(doc: Document, text: str, level: int) -> None:
    sizes = {2: 13, 3: 11.5, 4: 10.5}
    par = doc.add_paragraph()
    pf = par.paragraph_format
    pf.space_before = Pt(16 if level == 2 else 11)
    pf.space_after = Pt(5)
    pf.keep_with_next = True
    r = par.add_run(text)
    r.bold = True
    run_font(r, BODY_FONT, sizes.get(level, 10))
    if level == 2:
        # 항목 제목 아래 밑줄로 구획을 만든다
        pbdr = par._p.get_or_add_pPr()
        from docx.oxml import OxmlElement

        bd = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:color"), "999999")
        bd.append(bottom)
        pbdr.append(bd)


def bullet(doc: Document, text: str, depth: int) -> None:
    par = doc.add_paragraph()
    pf = par.paragraph_format
    pf.left_indent = Cm(0.5 + 0.5 * depth)
    pf.first_line_indent = Cm(-0.32)
    pf.space_after = Pt(2)
    r = par.add_run("· " if depth == 0 else "- ")
    run_font(r)
    add_inline(par, text)


def code_block(doc: Document, lines: list[str]) -> None:
    par = doc.add_paragraph()
    pf = par.paragraph_format
    pf.left_indent = Cm(0.3)
    pf.space_before = pf.space_after = Pt(6)
    pf.line_spacing = 1.15
    for i, l in enumerate(lines):
        if i:
            par.add_run("\n")
        r = par.add_run(l)
        run_font(r, MONO_FONT, 8.5)
    shade(par)


def shade(par) -> None:
    from docx.oxml import OxmlElement

    pPr = par._p.get_or_add_pPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), "F4F5F6")
    pPr.append(sh)


def table(doc: Document, rows: list[list[str]]) -> None:
    cols = max(len(r) for r in rows)
    t = doc.add_table(rows=0, cols=cols)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = True
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci in range(cols):
            cell = cells[ci]
            cell.paragraphs[0].paragraph_format.space_after = Pt(1)
            cell.paragraphs[0].paragraph_format.line_spacing = 1.2
            txt = row[ci] if ci < len(row) else ""
            add_inline(cell.paragraphs[0], txt, 9)
            if ri == 0:
                for r in cell.paragraphs[0].runs:
                    r.bold = True
                shade(cell.paragraphs[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def quote(doc: Document, text: str) -> None:
    par = doc.add_paragraph()
    pf = par.paragraph_format
    pf.left_indent = Cm(0.4)
    pf.space_before = pf.space_after = Pt(5)
    add_inline(par, text)
    for r in par.runs:
        r.font.color.rgb = MUTED
        r.italic = True


# ── 머리말: 제목 + 팀명·구성원 ─────────────────────────────────
def header(doc: Document, spec: dict) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(spec["tag"])
    run_font(r, BODY_FONT, 9)
    r.font.color.rgb = MUTED

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(spec["title"])
    r.bold = True
    run_font(r, BODY_FONT, 17)

    t = doc.add_table(rows=2, cols=2)
    t.style = "Table Grid"
    for i, (k, hint) in enumerate(
        [("팀명", "등록된 팀명과 동일하게 작성"), ("구성원 성명", "팀장, 팀원 순으로 작성")]
    ):
        kc, vc = t.rows[i].cells
        kc.width, vc.width = Cm(3.4), Cm(13.6)
        add_inline(kc.paragraphs[0], f"**{k}**", 10)
        shade(kc.paragraphs[0])
        val = TEAM.get(k, "").strip()
        par = vc.paragraphs[0]
        run = par.add_run(val if val else hint)
        run_font(run, BODY_FONT, 10)
        if not val:
            run.font.color.rgb = MUTED

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("( * 필수항목)")
    run_font(r, BODY_FONT, 9)
    r.font.color.rgb = MUTED


# ── 변환 ─────────────────────────────────────────────────────
def convert(spec: dict) -> tuple[str, int]:
    src = pathlib.Path(spec["src"])
    lines = src.read_text(encoding="utf-8").splitlines()

    doc = Document()
    setup(doc)
    header(doc, spec)

    i = 0
    started = False  # 첫 "## " 를 만나기 전 머리말·전달 메모는 버린다
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        stripped = line.strip()

        if stripped.startswith("## "):
            title = stripped[3:].strip()
            # 양식이 "미구현 또는 향후 구현 예정 기능은 제외" 를 명시한다.
            # 내부 확인용으로 남긴 부록은 제출본에 넣지 않는다.
            if any(k in title for k in ("부록", "구현하지 않은", "제출본에는")):
                break
            started = True
            heading(doc, title, 2)
            i += 1
            continue

        if not started:
            i += 1
            continue

        if stripped.startswith("#### "):
            heading(doc, stripped[5:].strip(), 4)
        elif stripped.startswith("### "):
            heading(doc, stripped[4:].strip(), 3)
        elif stripped.startswith("```"):
            buf = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            code_block(doc, buf)
        elif stripped.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{2,}:?", c) for c in cells if c):
                    rows.append(cells)
                i += 1
            if rows:
                table(doc, rows)
            continue
        elif stripped.startswith(">"):
            # 인용은 여러 줄에 걸친다. 줄 단위로 처리하면 두 줄에 걸친 **굵게** 가 끊긴다.
            buf = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                seg = lines[i].strip().lstrip(">").strip()
                if seg:
                    buf.append(seg)
                else:
                    if buf:
                        quote(doc, " ".join(buf))
                        buf = []
                i += 1
            if buf:
                quote(doc, " ".join(buf))
            continue
        elif re.match(r"^\s*[-*] ", raw):
            depth = (len(raw) - len(raw.lstrip())) // 2
            text = re.sub(r"^\s*[-*] ", "", raw)
            # 다음 줄이 이어지는 들여쓰기면 한 항목으로 합친다
            while i + 1 < len(lines):
                nxt = lines[i + 1]
                is_block = re.match(r"^\s*(?:[-*+]\s|\d+\.\s|#{1,6}\s|\||>\s|```)", nxt)
                if nxt.strip() and not is_block and nxt.startswith("  "):
                    text += " " + nxt.strip()
                    i += 1
                else:
                    break
            bullet(doc, text, depth)
        elif re.match(r"^\s*\d+\. ", raw):
            bullet(doc, re.sub(r"^\s*\d+\. ", "", raw), 0)
        elif stripped in ("", "---"):
            pass
        else:
            text = stripped
            while i + 1 < len(lines):
                nxt = lines[i + 1].strip()
                is_block = re.match(r"^(?:[-*+]\s|\d+\.\s|#{1,6}\s|\||>\s|```)", nxt)
                if nxt and not is_block and nxt != "---":
                    text += " " + nxt
                    i += 1
                else:
                    break
            par = doc.add_paragraph()
            par.paragraph_format.space_after = Pt(5)
            add_inline(par, text)

        i += 1

    out = pathlib.Path(spec["out"])
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return str(out), len(doc.element.body)


if __name__ == "__main__":
    for spec in DOCS:
        path, blocks = convert(spec)
        print(f"  {path}  ({blocks} 블록)")
    if not TEAM["팀명"]:
        print("\n  팀명·구성원이 비어 있습니다. tools/md-to-docx.py 의 TEAM 을 채우고 다시 실행하세요.")
